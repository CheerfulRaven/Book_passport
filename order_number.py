"""Функция new_order_number присваивает новый номер заказа, возвращает значение номера заказа, а так же записывает номер
заказа в файл 'Номера заказов_2023.docx'"""

from docx import Document
from docx.shared import Mm
from datetime import datetime as dt


def new_order_number(arg1, arg2='', arg3=''):
    book_name = arg1
    author = arg2
    book_type = arg3
    # начинаем работу с документом 'Номера заказов_2023.docx'
    order_numbers_doc = Document('Номера заказов_2023.docx')
    table_2 = order_numbers_doc.tables[0]
    order_number = ''
    for i, row in enumerate(table_2.rows):
        # если таблица закончилась, продлеваем ее
        if i == len(table_2.rows) - 1:
            cells = table_2.add_row().cells
            row.cells[0].text = ''
        # заполняем первую пустую строку
        if row.cells[0].text == '':
            now = dt.now()
            date = now.strftime("%d.%m.%Y")
            row.cells[0].text = date
            row.cells[1].text = author
            # разделитель
            sep = " " if book_name[-1] in [".", "!", "?"] else ". "
            row.cells[2].text = book_name + sep + book_type.capitalize()
            if i != 1:
                order_number = int(table_2.rows[i-1].cells[4].text) + 1
                order_number = f'{order_number:0>3}'
            else:
                order_number = f'{1:0>3}'
            row.cells[4].text = order_number
            # изменяем шрифт
            row.cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[2].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[4].paragraphs[0].runs[0].font.name = 'Times New Roman'
            # убираем абзацный отступ
            row.cells[0].paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            row.cells[1].paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            row.cells[2].paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            row.cells[4].paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            break

    # следим за тем, чтобы документ не был открыт в другой программе; сохраняем документ
    try:
        order_numbers_doc.save('Номера заказов_2023.docx')
    except PermissionError:
        return "Закройте документ 'Номера заказов_2023'!"

    return order_number