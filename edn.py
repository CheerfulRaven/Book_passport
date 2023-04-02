"""Функция new_edn присваивает EDN из списка, который находится в файле, возвращает значение EDN (в виде ссылки), а
так же записывает данные для выбранного EDN в файл 'EDN резерв.docx'"""

from docx import Document
from docx.shared import Mm


def new_edn(arg1, arg2='', arg3=''):
    book_name = arg1
    author = arg2
    book_type = arg3
    edn = ""
    i = 0
    # начинаем работу с документом 'EDN резерв.docx'
    edn_doc = Document('EDN резерв.docx')
    # получаем первую таблицу в документе
    table_3 = edn_doc.tables[0]
    # читаем данные из таблицы
    for row in table_3.rows:
        if row.cells[2].text == '':
            if author != "":
                # sep1 - разделитель
                sep1 = " " if author[-1] == "." else ". "
            else:
                sep1 = ""
            # sep2 - разделитель
            sep2 = " " if book_name[-1] in [".", "!", "?"] else ". "
            info = author + sep1 + book_name + sep2 + book_type.capitalize()
            row.cells[2].text = info
            # изменяем шрифт
            row.cells[2].paragraphs[0].runs[0].font.name = 'Times New Roman'
            # убираем абзацный отступ
            row.cells[2].paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            edn = row.cells[3].text
            # указываем порядковый номер
            row.cells[0].text = str(i) + "."
            break
        else:
            i += 1

    if edn == "":
        return "Добавьте EDN!"

    # следим за тем, чтобы документ не был открыт в другой программе; сохраняем документ
    try:
        edn_doc.save('EDN резерв.docx')
    except PermissionError:
        return "Закройте документ 'EDN резерв'!"

    return edn
