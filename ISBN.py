"""Функция new_isbn присваивает ISBN из списка, который находится в файле, возвращает значение ISBN, а так же записывает
данные для выбранного ISBN в файл 'ISBN_2023.docx'"""

from docx import Document
from docx.shared import Mm
from datetime import datetime as dt


def new_isbn(arg1, arg2='', arg3=''):
    book_name = arg1
    author = arg2
    book_type = arg3
    # info - информация о книге, date - дата присвоения isbn, isbn - ISBN, order_number - номер заказа, edn - код EDN
    info = ''
    result = ''
    isbn = ''
    # начинаем работу с документом 'ISBN_2023.docx'
    doc = Document('ISBN_2023.docx')
    # получаем первую таблицу в документе
    table = doc.tables[0]
    # читаем данные из таблицы
    for row in table.rows:
        info = row.cells[2]
        isbn = row.cells[1].text
        # заполняем первую пустую строку
        if row.cells[2].text == '':
            # исключаем ситуацию, когда у книги нет автора
            if author != "":
                # sep1 - разделитель
                sep1 = " " if author[-1] == "." else ". "
            else:
                sep1 = ""
            # sep2 - разделитель
            sep2 = " " if book_name[-1] in [".", "!", "?"] else ". "
            info = author + sep1 + book_name + sep2 + book_type.capitalize()
            row.cells[2].text = info
            now = dt.now()
            date = now.strftime("%d.%m.%Y")
            row.cells[4].text = date
            # изменяем шрифт
            row.cells[2].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[4].paragraphs[0].runs[0].font.name = 'Times New Roman'
            # убираем абзацный отступ
            row.cells[2].paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            row.cells[4].paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            # выводим данные для дальнейшей обработки
            result = isbn + ' ' + info
            break

    if result == "":
        return "Добавьте ISBN"

    # следим за тем, чтобы документ не был открыт в другой программе; сохраняем документ
    try:
        doc.save('ISBN_2023.docx')
    except PermissionError:
        return "Закройте документ 'ISBN_2023'!"

    return isbn